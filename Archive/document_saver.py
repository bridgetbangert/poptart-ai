import docx
from docx.shared import Cm

# Create a document
doc = docx.Document()
sections = doc.sections

for section in sections:
    section.top_margin = Cm(.5)
    section.bottom_margin = Cm(.5)
    section.left_margin = Cm(.5)
    section.right_margin = Cm(.5)

# Add a paragraph to the document
p = doc.add_paragraph()

# Add some formatting to the paragraph
p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0

# Add a run to the paragraph
run = p.add_run("python-docx")

# Add some formatting to the run
run.bold = True
run.italic = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Add more text to the same paragraph
run = p.add_run("Tutorial")

# Format the run
run.bold = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Add another paragraph (left blank for an empty line)
doc.add_paragraph()

# Add another paragraph
p = doc.add_paragraph()

# Add a run and format it
run = p.add_run("This is my first python-docx tutorial!")
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(12)